import os
import json
import io
from datetime import datetime

from dotenv import load_dotenv
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches

# ───── Load API key ─────────────────────────────────────────────────────────
load_dotenv()
GEMMA_KEY = os.getenv("GEMMA_API_KEY")
if not GEMMA_KEY:
    st.error("Missing GEMMA_API_KEY in .env")
    st.stop()
genai.configure(api_key=GEMMA_KEY)

# ───── Load parameters from JSON ─────────────────────────────────────────────
HERE      = os.path.dirname(__file__)
json_path = os.path.join(HERE, "recipace.json")
with open(json_path, "r", encoding="utf-8") as f:
    PARAMS = json.load(f)

# ───── Page setup ───────────────────────────────────────────────────────────
st.set_page_config(page_title="Recipace-AI", layout="wide", page_icon="🥣")

# Header with logo + title



st.markdown("<h1 style='margin-bottom:0'>🥣 Recipace-AI</h1>", unsafe_allow_html=True)
st.markdown("Generate richly detailed recipes, scripts & more for your YouTube cooking channel.")

st.write("---")

# ───── Session state defaults ────────────────────────────────────────────────
defaults = {
    "saved": [],
    "last_prompt": None,
    "last_recipe": None,
    "unit_system": "Imperial",
    "rating": 3
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v
for p in PARAMS:
    if p not in st.session_state:
        st.session_state[p] = "None"

# ───── Tabs ─────────────────────────────────────────────────────────────────
tab_gen, tab_saved, tab_about = st.tabs(["Generate", "Saved Recipes", "About"])

with tab_gen:
    # Toggles
    chef_mode    = st.checkbox("👩‍🍳 Chef Mode (plating & pairing tips)")
    video_script = st.checkbox("🎥 Include YouTube Video Script")
    regenerate   = st.button("🔄 Regenerate Last Recipe",
                             disabled=st.session_state.last_prompt is None)

    # Parameters accordion
    with st.expander("🎛️ Select Your Recipe Parameters", expanded=False):
        with st.form("params_form", clear_on_submit=False):
            cols = st.columns(3)
            for idx, (label, opts) in enumerate(PARAMS.items()):
                col = cols[idx % 3]
                default = st.session_state.get(label, "None")
                default_idx = opts.index(default) if default in opts else 0
                col.selectbox(label, opts, index=default_idx, key=label)
            generate = st.form_submit_button("🍳 Generate Recipe")

    if generate or regenerate:
        # Build prompt with extra detail request
        parts = [f"{lbl}: {st.session_state[lbl]}" 
                 for lbl in PARAMS if st.session_state[lbl] != "None"]
        prompt = (
            "Generate a one-of-a-kind recipe (title, ingredients, cooking method, step-by-step guide) "
            "for a YouTube food video creator with the following parameters:\n"
            + "\n".join(parts)
            + "\n\nInclude detailed explanations, chef’s tips, and common mistakes to avoid at each step."
        )
        if chef_mode:
            prompt += "\nAlso include professional plating suggestions and beverage pairings."
        if video_script:
            prompt += (
                "\nAfter the recipe, generate a concise YouTube video script "
                "(intro, step narration, outro)."
            )

        # Extras in one shot
        prompt += (
            "\n\nThen also produce, under clear headings:\n"
            "- Grocery List\n"
            "- Step-by-Step Timers (with estimated durations)\n"
            "- Nutrition Facts\n"
            "- Ingredient Substitutions"
        )

        st.session_state.last_prompt = prompt
        with st.spinner("🌟 Generating recipe with full detail…"):
            model    = genai.GenerativeModel("gemma-3-27b-it")
            response = model.generate_content(prompt)
        st.session_state.last_recipe = response.text

    if st.session_state.last_recipe:
        # Display area
        st.markdown("### 🍽️ Your AI-Generated Recipe & Extras")
        st.code(st.session_state.last_recipe, language=None)

        st.write("---")

        # Unit conversion
        units = st.radio("📏 Unit System", ["Imperial", "Metric"],
                         index=["Imperial","Metric"].index(st.session_state.unit_system),
                         key="unit_system")
        if st.button("📐 Convert Units"):
            with st.spinner("🔄 Converting units…"):
                uc = model.generate_content(
                    f"Convert measurements to {units} units:\n{st.session_state.last_recipe}"
                )
            st.markdown(f"#### Recipe in {units}")
            st.write(uc.text)

        # Rating & share
        st.slider("⭐ Rate this Recipe", 1, 5, key="rating")
        tweet = st.session_state.last_recipe.replace("\n", "%0A")[:250]
        st.markdown(f"[🐦 Share on Twitter]"
                    f"(https://twitter.com/intent/tweet?text={tweet})")

        # Save / Download buttons
        col_save, col_txt, col_doc = st.columns(3)
        with col_save:
            if st.button("💾 Save Recipe"):
                st.session_state.saved.append(st.session_state.last_recipe)
                st.success("Saved!")
        with col_txt:
            st.download_button("⬇️ Download .txt",
                st.session_state.last_recipe,
                file_name="recipace_recipe.txt",
                mime="text/plain")
        with col_doc:
            # Build detailed docx
            def make_docx(text: str, params: dict) -> io.BytesIO:
                doc = Document()
                # Title
                lines = text.split("\n")
                title = lines[0]
                doc.add_heading(title, level=0)
                doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                # Parameters summary
                doc.add_heading("Parameters", level=1)
                for k, v in params.items():
                    if v != "None":
                        doc.add_paragraph(f"• {k}: {v}")
                # Main content
                doc.add_heading("Recipe & Extras", level=1)
                for line in lines[1:]:
                    if line.strip().endswith(":"):
                        doc.add_heading(line.rstrip(":"), level=2)
                    else:
                        doc.add_paragraph(line)
                # Footer
                doc.add_page_break()
                doc.add_paragraph("Generated by Recipace-AI")
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf

            docx_buf = make_docx(st.session_state.last_recipe,
                                 {k: st.session_state[k] for k in PARAMS})
            st.download_button("📄 Download .docx",
                data=docx_buf,
                file_name="recipace_recipe.docx",
                mime=(
                    "application/"
                    "vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ))

with tab_saved:
    st.header("📚 Saved Recipes")
    if not st.session_state.saved:
        st.info("No saved recipes yet.")
    else:
        for i, rec in enumerate(st.session_state.saved, 1):
            with st.expander(f"Recipe {i}"):
                st.code(rec, language=None)
                st.download_button(
                    f"⬇️ Save {i} as .txt",
                    data=rec,
                    file_name=f"saved_{i}.txt",
                    mime="text/plain"
                )

with tab_about:
    st.header("ℹ️ About Recipace-AI")
    left, right = st.columns([3,1])
    with left:
        st.markdown("""
**Recipace-AI** is your all-in-one cooking companion for YouTube creators:

- 🔍 **30+ Customizable Parameters**: From cuisine style to spice levels.  
- 🍳 **Rich, Step-by-Step Guides**: Detailed explanations, chef tips & common mistakes.  
- 📋 **Extras Included**: Grocery lists, timers, nutrition facts & substitutions.  
- 🎥 **Optional Video Script**: Ready-to-use intro, narration & outro.  
- 🔄 **Scaling & Conversion**: Adjust servings & switch between Imperial/Metric.  
- 💾 **Save & Download**: Export as plain text or polished Word document.

Built with Streamlit + Google Gemma 3.27B.
    """)
    with right:
        st.image("logo.png", use_container_width=True)
    st.markdown("> **Get started** on the **Generate** tab, select your parameters, and hit **Generate Recipe**!")
